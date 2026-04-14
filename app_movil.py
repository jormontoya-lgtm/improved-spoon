import streamlit as st
import sqlite3
import pandas as pd
from io import BytesIO
from datetime import datetime, timedelta
import base64
from docx import Document

# --- CONFIGURACIÓN ---
USUARIOS_PERMITIDOS = {"jorge": "1234", "supervisor": "obra2026", "Gerardo": "123456", "Julie": "789011", "Diego": "121314"}
STOCK_MINIMO = 20 

st.set_page_config(page_title="SGO-H Pro", layout="centered")

# --- FUNCIONES DE BASE DE DATOS ---
def conectar():
    conn = sqlite3.connect("sistema_obra.db")
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS reportes 
                   (id INTEGER PRIMARY KEY AUTOINCREMENT, fecha TEXT, operador TEXT, 
                    tramo TEXT, actividad TEXT, material TEXT, avance REAL, 
                    observaciones TEXT, fotos TEXT, editado TEXT)''')
    cur.execute('''CREATE TABLE IF NOT EXISTS entradas_almacen 
                   (id INTEGER PRIMARY KEY AUTOINCREMENT, fecha TEXT, material TEXT, cantidad REAL, 
                    autoriza TEXT, verificado TEXT)''')
    cur.execute('''CREATE TABLE IF NOT EXISTS inventario 
                   (id INTEGER PRIMARY KEY AUTOINCREMENT, material TEXT, cantidad REAL)''')
    cur.execute('''CREATE TABLE IF NOT EXISTS logs 
                   (id INTEGER PRIMARY KEY AUTOINCREMENT, fecha TEXT, usuario TEXT, accion TEXT)''')
    
    cur.execute("SELECT COUNT(*) FROM inventario")
    if cur.fetchone()[0] == 0:
        mats = [('Tubo PVC 2"', 100), ('Tubo PVC 4"', 100), ('Tubo PVC 6"', 100), 
                ('Tubo PVC 8"', 100), ('Tubo PVC 10"', 100), ('Tubo PVC 12"', 100),
                ('Cemento (Sacos)', 100), ('Varilla 1/2', 100)]
        cur.executemany("INSERT INTO inventario (material, cantidad) VALUES (?,?)", mats)
    conn.commit()
    return conn

def registrar_log(usuario, accion):
    conn = conectar(); cur = conn.cursor()
    cur.execute("INSERT INTO logs (fecha, usuario, accion) VALUES (?,?,?)", 
                (obtener_hora_local().strftime("%Y-%m-%d %H:%M:%S"), usuario, accion))
    conn.commit(); conn.close()

def obtener_hora_local():
    return (datetime.utcnow() - timedelta(hours=6))

# --- AUTENTICACIÓN ---
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

if not st.session_state.autenticado:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        try: st.image("logo.png", width=250)
        except: st.info("SGO-H")
    st.title("🚧 Acceso SGO-H")
    u = st.text_input("Usuario", key="u_login")
    p = st.text_input("Contraseña", type="password", key="p_login")
    if st.button("Entrar", use_container_width=True):
        if u.lower() in USUARIOS_PERMITIDOS and USUARIOS_PERMITIDOS[u.lower()] == p:
            st.session_state.autenticado = True
            st.session_state.usuario_actual = u.lower()
            registrar_log(u, "Inicio de Sesión")
            st.rerun()
        else: st.error("Credenciales incorrectas")

else:
    # Sidebar
    try: st.sidebar.image("logo.png", width=100)
    except: pass
    st.sidebar.title(f"👤 {st.session_state.usuario_actual.capitalize()}")
    menu_ops = ["Reportar Avance", "Entrada Almacén", "Ver Inventario", "Exportar"]
    if st.session_state.usuario_actual == "jorge":
        menu_ops.insert(0, "Panel de Control Jorge")
    menu = st.sidebar.selectbox("Ir a:", menu_ops)
    st.sidebar.divider()
    if st.sidebar.button("🔴 Cerrar Sesión", use_container_width=True):
        st.session_state.autenticado = False
        st.rerun()
    if st.session_state.usuario_actual == "jorge":
        if st.sidebar.button("🗑️ RESETEAR TODO", use_container_width=True):
            conn = conectar(); cur = conn.cursor()
            cur.execute("DROP TABLE IF EXISTS reportes"); cur.execute("DROP TABLE IF EXISTS inventario")
            cur.execute("DROP TABLE IF EXISTS entradas_almacen"); cur.execute("DROP TABLE IF EXISTS logs")
            conn.commit(); conn.close(); st.rerun()

    # --- SECCIONES ---
    if menu == "Panel de Control Jorge":
        st.header("📋 Historial Administrativo")
        conn = conectar()
        df = pd.read_sql_query("SELECT fecha, operador, tramo, actividad, material, avance, observaciones FROM reportes ORDER BY fecha DESC", conn)
        conn.close()
        st.dataframe(df, use_container_width=True)

    elif menu == "Reportar Avance":
        st.header("📝 Nuevo Reporte de Obra")
        
        # --- LEYENDA DE ENVIADO ---
        if st.session_state.get('rep_listo', False):
            st.success("✅ Reporte enviado, gracias por tu compromiso")
            if st.button("Hacer otro reporte"):
                st.session_state.rep_listo = False
                st.rerun()
        else:
            # --- FORMULARIO DE REPORTE ---
            tra = st.text_input("Tramo / Ubicación")
            act = st.selectbox("Actividad", ["Excavación", "Instalación de Tubería", "Relleno", "Armado"])
            mat_f = "N/A"
            if act == "Instalación de Tubería":
                mat_f = st.selectbox("Diámetro:", ['Tubo PVC 2"', 'Tubo PVC 4"', 'Tubo PVC 6"', 'Tubo PVC 8"', 'Tubo PVC 10"', 'Tubo PVC 12"'])
            elif act == "Relleno": mat_f = "Cemento (Sacos)"
            elif act == "Armado": mat_f = "Varilla 1/2"
            
            ava = st.number_input("Cantidad / Metros:", min_value=0.0)
            obs = st.text_area("🗒️ Observaciones", placeholder="Detalles imprevistos...")
            archivos = st.file_uploader("📸 Fotos", accept_multiple_files=True, type=['png', 'jpg', 'jpeg'])
            
            if st.button("💾 GUARDAR REPORTE", use_container_width=True):
                if tra and ava >= 0:
                    f_str = "|".join([base64.b64encode(a.getvalue()).decode() for a in archivos[:3]])
                    conn = conectar(); cur = conn.cursor()
                    cur.execute("""INSERT INTO reportes 
                                   (fecha, operador, tramo, actividad, material, avance, observaciones, fotos) 
                                   VALUES (?,?,?,?,?,?,?,?)""",
                                (obtener_hora_local().strftime("%Y-%m-%d %H:%M"), 
                                 st.session_state.usuario_actual, tra, act, mat_f, ava, obs, f_str))
                    
                    if mat_f != "N/A":
                        cur.execute("UPDATE inventario SET cantidad = cantidad - ? WHERE material = ?", (ava, mat_f))
                    
                    conn.commit(); conn.close()
                    st.session_state.rep_listo = True
                    st.rerun()
                else:
                    st.warning("Por favor, completa al menos el tramo.")

        # --- SECCIÓN: VER MIS REPORTES (HISTORIAL PERSONAL CON FOTOS) ---
        st.divider()
        if st.button("🔍 Ver mis reportes enviados", use_container_width=True):
            st.session_state.ver_historial_personal = not st.session_state.get('ver_historial_personal', False)

        if st.session_state.get('ver_historial_personal', False):
            st.subheader(f"📂 Mis Reportes - {st.session_state.usuario_actual.capitalize()}")
            conn = conectar()
            # Traemos también la columna 'fotos'
            query = "SELECT fecha, tramo, actividad, material, avance, observaciones, fotos FROM reportes WHERE operador = ? ORDER BY id DESC"
            df_personal = pd.read_sql_query(query, conn, params=(st.session_state.usuario_actual,))
            conn.close()

            if not df_personal.empty:
                # Mostramos la tabla (sin la columna de fotos cruda porque es texto largo)
                st.dataframe(df_personal.drop(columns=['fotos']), use_container_width=True)
                
                # --- GALERÍA DE FOTOS ABAJO DE LA TABLA ---
                st.write("### 📸 Galería de fotos por reporte:")
                for index, row in df_personal.iterrows():
                    if row['fotos']:
                        with st.expander(f"Fotos del {row['fecha']} - Tramo: {row['tramo']}"):
                            # Separamos las fotos (están unidas por '|')
                            lista_fotos = row['fotos'].split('|')
                            cols_fotos = st.columns(len(lista_fotos))
                            for i, foto_b64 in enumerate(lista_fotos):
                                if foto_b64:
                                    img_data = base64.b64decode(foto_b64)
                                    cols_fotos[i].image(img_data, use_container_width=True)
                
                # Botón para descargar el Excel (el Excel no guarda imágenes, solo texto)
                buf_p = BytesIO()
                with pd.ExcelWriter(buf_p, engine='openpyxl') as wr:
                    df_personal.drop(columns=['fotos']).to_excel(wr, index=False, sheet_name='Mis_Reportes')
                
                st.download_button(
                    label="📥 Descargar mi historial en Excel",
                    data=buf_p.getvalue(),
                    file_name=f"Mis_Reportes_{st.session_state.usuario_actual}.xlsx",
                    use_container_width=True
                )
            else:
                st.info("Aún no has enviado ningún reporte.")

    elif menu == "Entrada Almacén":
        st.header("📥 Entrada Almacén")
        if st.session_state.get('ent_listo', False):
            st.success("✅ Entrada registrada, gracias por tu compromiso")
            if st.button("Registrar otra entrada"): st.session_state.ent_listo = False; st.rerun()
        else:
            with st.form("ent_f"):
                m_e = st.selectbox("Material:", ['Tubo PVC 2"', 'Tubo PVC 4"', 'Tubo PVC 6"', 'Tubo PVC 8"', 'Tubo PVC 10"', 'Tubo PVC 12"', 'Cemento (Sacos)', 'Varilla 1/2'])
                c_e = st.number_input("Cantidad:", min_value=0.1)
                aut = st.text_input("Autoriza:")
                if st.form_submit_button("REGISTRAR ENTRADA"):
                    conn = conectar(); cur = conn.cursor()
                    cur.execute("INSERT INTO entradas_almacen (fecha, material, cantidad, autoriza) VALUES (?,?,?,?)",
                                (obtener_hora_local().strftime("%Y-%m-%d %H:%M"), m_e, c_e, aut))
                    cur.execute("UPDATE inventario SET cantidad = cantidad + ? WHERE material = ?", (c_e, m_e))
                    conn.commit(); conn.close(); st.session_state.ent_listo = True; st.rerun()

    elif menu == "Ver Inventario":
        st.header("📦 Stock Actual")
        conn = conectar(); df_i = pd.read_sql_query("SELECT material, cantidad FROM inventario", conn); conn.close()
        if (df_i['cantidad'] < STOCK_MINIMO).any():
            st.error("⚠️ STOCK BAJO")
            st.markdown('<audio src="https://www.soundjay.com/buttons/beep-01a.mp3" autoplay></audio>', unsafe_allow_html=True)
        def color_inv(v): return 'background-color: #ff4b4b; color: white' if v < STOCK_MINIMO else ''
        st.dataframe(df_i.style.map(color_inv, subset=['cantidad']), use_container_width=True)

    elif menu == "Exportar":
        st.header("📊 Gestión de Informes")
        if st.button("🚀 ENVIAR INFORME", use_container_width=True):
            registrar_log(st.session_state.usuario_actual, "Informe Generado")
            st.session_state.exp_listo = True
            st.success("Reporte procesado.")

        if st.session_state.get('exp_listo', False):
            conn = conectar()
            df_r = pd.read_sql_query("SELECT * FROM reportes", conn)
            df_stock = pd.read_sql_query("SELECT * FROM inventario", conn)
            df_ent = pd.read_sql_query("SELECT * FROM entradas_almacen", conn)
            fecha_f = obtener_hora_local().strftime("%Y-%m-%d")

            # EXCEL MULTI-HOJA
            buffer = BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as wr:
                df_r.drop(columns=['fotos']).to_excel(wr, index=False, sheet_name='Bitacora_Obra')
                df_stock.to_excel(wr, index=False, sheet_name='Stock_Actual')
                df_ent.to_excel(wr, index=False, sheet_name='Historial_Entradas')
            
            st.download_button(
                label="📥 Descargar Excel Completo",
                data=buffer.getvalue(),
                file_name=f"Reporte_Obra_{fecha_f}.xlsx",
                use_container_width=True
            )

            if st.session_state.usuario_actual == "jorge":
                df_l = pd.read_sql_query("SELECT * FROM logs ORDER BY id DESC", conn)
                doc = Document(); doc.add_heading(f'Auditoría de Sistema - {fecha_f}', 0)
                for _, r in df_l.iterrows(): doc.add_paragraph(f"{r['fecha']} - {r['usuario']}: {r['accion']}")
                b_w = BytesIO(); doc.save(b_w)
                st.download_button("📄 Descargar Word (Auditoría)", b_w.getvalue(), f"Auditoria_{fecha_f}.docx", use_container_width=True)
            conn.close()